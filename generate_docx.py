import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def build_docx(filename="FarmBuddy_PRD.docx"):
    doc = docx.Document()
    
    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Styles helper
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(60, 60, 60)
    
    # Document Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Farm Buddy PRD")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(26)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(20, 40, 30)
    
    # Add spacing after title
    title.paragraph_format.space_after = Pt(24)
    
    # Section Heading Helper
    def add_heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(16, 124, 65) # Emerald Green
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        return p

    # Paragraph Helper
    def add_body_text(text):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        return p
        
    # List Item Helper
    def add_list_item(text, is_bullet=True):
        style = 'List Bullet' if is_bullet else 'List Bullet 2'
        p = doc.add_paragraph(style=style)
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(60, 60, 60)
        p.paragraph_format.space_after = Pt(3)
        return p

    # 1. Product Overview
    add_heading("1. Product Overview")
    add_body_text("Farm Buddy is a smart agriculture platform designed to help farmers monitor crops, maintain farm records, and provide supply chain traceability using AI and blockchain.")
    
    # 2. Problem Statement
    add_heading("2. Problem Statement")
    add_body_text("Farmers face crop diseases, lack of proper records, no supply chain transparency, and consumers cannot verify product authenticity.")
    
    # 3. Objectives
    add_heading("3. Objectives")
    add_list_item("AI-based crop guidance")
    add_list_item("Digital farm records")
    add_list_item("Disease detection")
    add_list_item("Supply chain traceability")
    add_list_item("Consumer trust")
    
    # 4. Target Users
    add_heading("4. Target Users")
    add_body_text("Primary: Farmers\nSecondary: Buyers, Consumers, Distributors")
    
    # 5. Core Features
    add_heading("5. Core Features")
    add_list_item("Farmer Dashboard")
    add_list_item("Crop Health Monitoring")
    add_list_item("AI Chatbot (Gemini API)")
    add_list_item("Activity Logging")
    add_list_item("QR Code Generation")
    add_list_item("Supply Chain Traceability")
    
    # 6. Technology Stack
    add_heading("6. Technology Stack")
    add_body_text("Frontend: React.js, Tailwind CSS\nBackend: Node.js, Express.js\nDatabase: MySQL\nAI: Gemini API\nBlockchain: Solidity (Future Phase)")
    
    # 7. System Flow
    add_heading("7. System Flow")
    flow_steps = [
        "1. Farmer Actions: Planting / Irrigation / Fertilization / Harvest / Lab Test etc.",
        "2. Data Ingestion Layer: Web/Mobile App, Farm Buddy Voice App, IoT Sensors.",
        "3. Backend API (Node.js/Express): Validates data, assigns Batch ID and timestamps.",
        "4. MySQL Database: Stores complete batch, event, operator, and sensor details.",
        "5. Hashing Engine: Generates SHA-256 hash of event data (Chain of Custody).",
        "6. File Storage (IPFS): Uploads lab reports/images, receives IPFS CID/URI.",
        "7. Smart Contract: recordFarmEvent(batchId, ipfsHash, dataHash, timestamp) on-chain.",
        "8. Blockchain Ledger: Stores Batch ID, Data Hash, IPFS CID, and Transaction ID.",
        "9. Batch Ready for Market: Quality checked and approved.",
        "10. QR Code Generation: Generates unique QR code with Batch ID and tracking URL.",
        "11. QR Code Printed on Packaging.",
        "12. Consumer Scans QR Code: Scans using any smartphone.",
        "13. Consumer Portal (Web App): Queries Blockchain and Backend API for Batch ID.",
        "14. Display Traceability Timeline: Displays crop events, sensors, and IPFS reports.",
        "15. Verification Process: Recalculates hash and fetches blockchain hash.",
        "16. Hash Match Verification: Compares recalculated hash with on-chain hash.",
        "17A. Verified: Hashes match. Data is authentic and untampered.",
        "17B. Tampered: Hashes do not match. Data is flagged as modified or corrupted."
    ]
    for step in flow_steps:
        add_body_text(step)
    
    # Page Break for clean 2-page print
    doc.add_page_break()
    
    # 8. Functional Requirements
    add_heading("8. Functional Requirements")
    add_body_text("FR1: User Registration/Login\nFR2: Add Farm Details\nFR3: Upload Crop Images\nFR4: AI Analysis\nFR5: Store Activity Logs\nFR6: Generate QR Codes")
    
    # 9. Non-Functional Requirements
    add_heading("9. Non-Functional Requirements")
    add_body_text("Fast response, secure storage, scalability, simple UI, multilingual support")
    
    # 10. Future Scope
    add_heading("10. Future Scope")
    add_body_text("IoT integration, weather forecasting, smart irrigation, blockchain verification, market prediction")
    
    # Conclusion
    add_heading("Conclusion")
    add_body_text("Farm Buddy aims to modernize agriculture through AI-powered assistance and transparent product traceability.")
    
    # Save document
    doc.save(filename)
    print(f"[Success] DOCX PRD successfully written to: {filename}")

if __name__ == "__main__":
    build_docx()
